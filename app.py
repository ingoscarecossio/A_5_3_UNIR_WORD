"""
Aplicación Streamlit Profesional para Combinar Documentos Word
Versión: 3.0 - Solución Robusta con docxcompose
Usa docxcompose para combinación profesional sin páginas en blanco
"""

import os
import logging
from io import BytesIO
from typing import List, Tuple, Dict, Optional
from datetime import datetime
import traceback

import streamlit as st
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuración de página
st.set_page_config(
    page_title="Combinador Profesional de Documentos Word",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================================
# ESTILOS CSS PROFESIONALES
# ============================================================================

st.markdown("""
    <style>
    .main {
        padding-top: 2rem;
    }
    
    .document-card {
        padding: 1.2rem;
        border: 2px solid #e0e0e0;
        border-radius: 12px;
        margin: 0.75rem 0;
        background: linear-gradient(135deg, #f5f7fa 0%, #ffffff 100%);
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    
    .document-card:hover {
        border-color: #4CAF50;
        box-shadow: 0 4px 12px rgba(76, 175, 80, 0.2);
        transform: translateY(-2px);
    }
    
    .stButton>button {
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        transform: scale(1.02);
    }
    
    h1 {
        color: #2c3e50;
        font-weight: 700;
    }
    
    h2 {
        color: #34495e;
        font-weight: 600;
        border-bottom: 3px solid #4CAF50;
        padding-bottom: 0.5rem;
    }
    </style>
""", unsafe_allow_html=True)

# ============================================================================
# CLASES PROFESIONALES
# ============================================================================

class DocumentInfo:
    """Clase para almacenar información de un documento"""
    def __init__(self, name: str, source_type: str, source, size: float = 0):
        self.name = name
        self.source_type = source_type
        self.source = source
        self.size = size
        self.paragraphs = 0
        self.tables = 0
        self.is_valid = True
        self.error_message = None
        self._analyzed = False
    
    def analyze(self):
        """Analiza el documento para obtener información detallada"""
        if self._analyzed:
            return
        
        try:
            if self.source_type == "path":
                doc = Document(self.source)
            else:
                self.source.seek(0)
                doc = Document(self.source)
                self.source.seek(0)
            
            self.paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            self.tables = len(doc.tables)
            self.is_valid = True
            self._analyzed = True
            
        except Exception as e:
            self.is_valid = False
            self.error_message = str(e)
            logger.error(f"Error analizando {self.name}: {e}")
            self._analyzed = True


class ProfessionalDocumentMerger:
    """Clase profesional para combinar documentos usando docxcompose"""
    
    def __init__(self, progress_callback=None):
        self.progress_callback = progress_callback
        self.stats = {
            'total_docs': 0,
            'total_paragraphs': 0,
            'total_tables': 0,
            'processing_time': 0
        }
    
    def _update_progress(self, current: int, total: int, message: str = ""):
        """Actualiza la barra de progreso"""
        if self.progress_callback:
            self.progress_callback(current, total, message)
    
    def _add_cover_page(self, doc: Document, options: Dict):
        """Agrega una portada profesional al documento"""
        # Insertar al inicio
        if len(doc.paragraphs) > 0:
            doc.paragraphs[0].insert_paragraph_before()
        
        # Usar add_paragraph en lugar de add_heading para evitar problemas con estilos
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(options.get('cover_title', 'Documentos Combinados'))
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar espacio
        doc.add_paragraph()
        
        subtitle_text = options.get('cover_subtitle', '') or f'Generado el {datetime.now().strftime("%d/%m/%Y %H:%M")}'
        if subtitle_text:
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(subtitle_text)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if options.get('cover_info', ''):
            doc.add_paragraph()  # Espacio
            info = doc.add_paragraph()
            info_run = info.add_run(options['cover_info'])
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_table_of_contents(self, doc: Document, documents: List[DocumentInfo]):
        """Agrega un índice de contenidos"""
        doc.add_page_break()
        
        # Usar párrafo en lugar de heading para evitar problemas con estilos
        toc_heading = doc.add_paragraph()
        toc_run = toc_heading.add_run('Índice de Contenidos')
        toc_run.bold = True
        toc_run.font.size = Pt(18)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        for idx, doc_info in enumerate(documents, 1):
            para = doc.add_paragraph()
            para.add_run(f"{idx}. {doc_info.name}")
    
    def merge_documents(
        self,
        documents: List[DocumentInfo],
        options: Dict
    ) -> Tuple[bytes, Dict]:
        """
        Combina múltiples documentos: Documento 1 completo, siguiente página, Documento 2 completo, etc.
        
        Args:
            documents: Lista de DocumentInfo ordenados
            options: Diccionario con opciones de combinación
        
        Returns:
            Tupla con (bytes del documento, estadísticas)
        """
        start_time = datetime.now()
        
        if not documents:
            raise ValueError("No hay documentos para combinar")
        
        self.stats['total_docs'] = len(documents)
        
        try:
            # Cargar TODOS los documentos primero para verificar que estén bien
            loaded_docs = []
            
            for idx, doc_info in enumerate(documents):
                self._update_progress(
                    idx,
                    len(documents),
                    f"Cargando: {doc_info.name}..."
                )
                
                try:
                    if doc_info.source_type == "path":
                        doc = Document(doc_info.source)
                    else:
                        doc_info.source.seek(0)
                        doc = Document(doc_info.source)
                        doc_info.source.seek(0)
                    
                    loaded_docs.append((doc_info, doc))
                    
                    # Actualizar estadísticas
                    self.stats['total_paragraphs'] += len([p for p in doc.paragraphs if p.text.strip()])
                    self.stats['total_tables'] += len(doc.tables)
                    
                except Exception as e:
                    logger.error(f"Error cargando {doc_info.name}: {e}")
                    if options.get('stop_on_error', False):
                        raise
                    continue
            
            if not loaded_docs:
                raise ValueError("No se pudieron cargar documentos válidos")
            
            # El primer documento es la base
            first_doc_info, master_doc = loaded_docs[0]
            
            # Agregar portada si está habilitado (ANTES de crear el compositor)
            if options.get('add_cover_page', False):
                self._add_cover_page(master_doc, options)
                if options.get('add_page_break', True):
                    master_doc.add_page_break()
            
            # Crear el compositor con el documento maestro
            composer = Composer(master_doc)
            
            # Procesar documentos restantes (desde el segundo en adelante)
            for idx, (doc_info, source_doc) in enumerate(loaded_docs[1:], start=2):
                self._update_progress(
                    idx - 1,
                    len(documents),
                    f"Combinando: {doc_info.name}..."
                )
                
                try:
                    # Si está habilitado el salto de página, cada documento va en nueva página
                    if options.get('add_page_break', False):
                        # Agregar documento completo en nueva página
                        composer.append(source_doc, break_type='page')
                    else:
                        # Agregar documento completo sin salto de página
                        composer.append(source_doc)
                    
                except Exception as e:
                    logger.error(f"Error combinando {doc_info.name}: {e}")
                    if options.get('stop_on_error', False):
                        raise
                    continue
            
            # Agregar índice si está habilitado (después de combinar todos)
            if options.get('add_table_of_contents', False):
                # Necesitamos agregar el índice al documento final
                # Guardamos temporalmente, agregamos índice, y volvemos a guardar
                temp_output = BytesIO()
                composer.save(temp_output)
                temp_output.seek(0)
                
                # Cargar el documento combinado
                final_doc = Document(temp_output)
                self._add_table_of_contents(final_doc, documents)
                
                # Guardar el documento final con índice
                output = BytesIO()
                final_doc.save(output)
                output.seek(0)
                result_bytes = output.read()
            else:
                # Guardar en memoria usando el compositor
                self._update_progress(len(documents), len(documents), "Guardando documento final...")
                output = BytesIO()
                composer.save(output)
                output.seek(0)
                result_bytes = output.read()
            
            # Calcular tiempo de procesamiento
            end_time = datetime.now()
            self.stats['processing_time'] = (end_time - start_time).total_seconds()
            
            return result_bytes, self.stats
            
        except Exception as e:
            logger.error(f"Error en merge_documents: {e}")
            logger.error(traceback.format_exc())
            raise

# ============================================================================
# FUNCIONES DE UTILIDAD
# ============================================================================

def format_file_size(size_bytes: float) -> str:
    """Formatea el tamaño del archivo en formato legible"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.2f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.2f} TB"

def get_file_size(file_path_or_obj) -> float:
    """Obtiene el tamaño del archivo en bytes"""
    try:
        if isinstance(file_path_or_obj, str):
            return os.path.getsize(file_path_or_obj)
        else:
            file_path_or_obj.seek(0, 2)
            size = file_path_or_obj.tell()
            file_path_or_obj.seek(0)
            return size
    except:
        return 0

def validate_docx_file(file_path_or_obj) -> Tuple[bool, Optional[str]]:
    """Valida que un archivo sea un .docx válido"""
    try:
        if isinstance(file_path_or_obj, str):
            doc = Document(file_path_or_obj)
        else:
            file_path_or_obj.seek(0)
            doc = Document(file_path_or_obj)
            file_path_or_obj.seek(0)
        return True, None
    except Exception as e:
        return False, str(e)

def list_docx_in_folder(folder_path: str) -> List[str]:
    """Lista todos los archivos .docx válidos en una carpeta"""
    if not folder_path or not os.path.isdir(folder_path):
        return []
    
    files = []
    for name in os.listdir(folder_path):
        full_path = os.path.join(folder_path, name)
        if (name.lower().endswith(".docx") and 
            not name.startswith("~$") and 
            os.path.isfile(full_path)):
            files.append(full_path)
    
    return sorted(files)

# ============================================================================
# INICIALIZACIÓN DE SESIÓN
# ============================================================================

if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'doc_sources' not in st.session_state:
    st.session_state.doc_sources = {}
if 'merged_bytes' not in st.session_state:
    st.session_state.merged_bytes = None
if 'merge_stats' not in st.session_state:
    st.session_state.merge_stats = None

# ============================================================================
# INTERFAZ DE USUARIO
# ============================================================================

# Título principal
st.title("📄 Combinador Profesional de Documentos Word")
st.markdown("**Versión 3.0 - Solución Robusta** | Usa docxcompose para combinación profesional sin páginas en blanco")

# Sidebar con opciones avanzadas
with st.sidebar:
    st.header("⚙️ Configuración")
    
    st.subheader("📋 Opciones de Combinación")
    add_page_breaks = st.checkbox("Agregar salto de página entre documentos", value=True)
    preserve_styles = st.checkbox("Preservar estilos originales", value=True, disabled=True, 
                                  help="Siempre activo con docxcompose")
    
    st.divider()
    
    st.subheader("📑 Elementos Adicionales")
    add_cover_page = st.checkbox("Agregar portada", value=False)
    if add_cover_page:
        cover_title = st.text_input("Título de portada", value="Documentos Combinados")
        cover_subtitle = st.text_input("Subtítulo", value="")
        cover_info = st.text_area("Información adicional", value="")
    
    add_table_of_contents = st.checkbox("Agregar índice de contenidos", value=False)
    
    st.divider()
    
    st.subheader("🔧 Opciones Avanzadas")
    stop_on_error = st.checkbox("Detener en caso de error", value=False)
    auto_analyze = st.checkbox("Analizar documentos automáticamente", value=True)
    
    st.divider()
    
    st.markdown("### 📊 Estadísticas")
    if st.session_state.merge_stats:
        stats = st.session_state.merge_stats
        st.metric("Documentos procesados", stats.get('total_docs', 0))
        st.metric("Párrafos totales", stats.get('total_paragraphs', 0))
        st.metric("Tablas totales", stats.get('total_tables', 0))
        st.metric("Tiempo de procesamiento", f"{stats.get('processing_time', 0):.2f}s")

# Sección de carga de documentos
st.header("📂 Cargar Documentos")

col1, col2 = st.columns([1, 1])

with col1:
    mode = st.radio(
        "Método de carga:",
        ["📁 Desde carpeta (local)", "📤 Subir archivos"],
        horizontal=False
    )

with col2:
    if mode == "📁 Desde carpeta (local)":
        st.info("ℹ️ **Nota**: Esta opción solo funciona en ejecución local. En Streamlit Cloud usa 'Subir archivos'.")
        folder = st.text_input(
            "Ruta de la carpeta",
            placeholder=r"C:\Users\...\Documentos",
            help="Ingresa la ruta completa de la carpeta (solo funciona localmente)"
        )
    else:
        uploaded = st.file_uploader(
            "Selecciona archivos .docx",
            type=["docx"],
            accept_multiple_files=True,
            help="Puedes seleccionar múltiples archivos"
        )

# Procesar carga de documentos
docs_info = []
doc_sources = {}

if mode == "📁 Desde carpeta (local)":
    if folder:
        paths = list_docx_in_folder(folder)
        if not paths:
            st.warning("⚠️ No se encontraron archivos .docx en esa carpeta")
        else:
            with st.spinner("Validando archivos..."):
                for path in paths:
                    name = os.path.basename(path)
                    is_valid, error = validate_docx_file(path)
                    
                    if is_valid:
                        size = get_file_size(path)
                        doc_info = DocumentInfo(name, "path", path, size)
                        
                        if auto_analyze:
                            doc_info.analyze()
                        
                        docs_info.append(doc_info)
                        doc_sources[name] = ("path", path)
                    else:
                        st.warning(f"⚠️ Archivo inválido: {name} - {error}")
            
            if docs_info:
                st.success(f"✅ {len(docs_info)} archivo(s) válido(s) cargado(s)")
else:
    if uploaded:
        with st.spinner("Validando archivos..."):
            for f in uploaded:
                is_valid, error = validate_docx_file(f)
                
                if is_valid:
                    size = get_file_size(f)
                    doc_info = DocumentInfo(f.name, "upload", f, size)
                    
                    if auto_analyze:
                        doc_info.analyze()
                    
                    docs_info.append(doc_info)
                    doc_sources[f.name] = ("upload", f)
                else:
                    st.warning(f"⚠️ Archivo inválido: {f.name} - {error}")
        
        if docs_info:
            st.success(f"✅ {len(docs_info)} archivo(s) válido(s) cargado(s)")

# Actualizar estado de sesión
st.session_state.documents = docs_info
st.session_state.doc_sources = doc_sources

if not docs_info:
    st.info("👆 Por favor, carga algunos documentos para comenzar")
    st.stop()

# Sección de visualización
st.header("📋 Documentos Cargados")

# Mostrar resumen
col1, col2, col3, col4 = st.columns(4)
total_size = sum(d.size for d in docs_info)
total_paragraphs = sum(d.paragraphs for d in docs_info)
total_tables = sum(d.tables for d in docs_info)

with col1:
    st.metric("Total Documentos", len(docs_info))
with col2:
    st.metric("Tamaño Total", format_file_size(total_size))
with col3:
    st.metric("Total Párrafos", total_paragraphs)
with col4:
    st.metric("Total Tablas", total_tables)

# Lista de documentos con controles de reordenamiento
st.subheader("🔄 Reordenar Documentos")

def move_up(index):
    if index > 0:
        docs_info[index], docs_info[index - 1] = docs_info[index - 1], docs_info[index]
        st.session_state.documents = docs_info
        st.rerun()

def move_down(index):
    if index < len(docs_info) - 1:
        docs_info[index], docs_info[index + 1] = docs_info[index + 1], docs_info[index]
        st.session_state.documents = docs_info
        st.rerun()

def remove_doc(index):
    doc_name = docs_info[index].name
    docs_info.pop(index)
    if doc_name in doc_sources:
        del doc_sources[doc_name]
    st.session_state.documents = docs_info
    st.session_state.doc_sources = doc_sources
    st.rerun()

# Mostrar lista de documentos
for idx, doc_info in enumerate(docs_info):
    with st.container():
        col1, col2, col3, col4, col5 = st.columns([0.4, 0.4, 5, 1.5, 0.7])
        
        with col1:
            if st.button("↑", key=f"up_{idx}", disabled=(idx == 0), use_container_width=True):
                move_up(idx)
        
        with col2:
            if st.button("↓", key=f"down_{idx}", disabled=(idx == len(docs_info) - 1), use_container_width=True):
                move_down(idx)
        
        with col3:
            st.markdown(f"**{idx + 1}.** {doc_info.name}")
            if doc_info._analyzed:
                st.caption(f"📄 {doc_info.paragraphs} párrafos | 📊 {doc_info.tables} tablas")
        
        with col4:
            st.caption(f"📦 {format_file_size(doc_info.size)}")
        
        with col5:
            if st.button("❌", key=f"remove_{idx}", use_container_width=True):
                remove_doc(idx)
        
        st.divider()

# Vista previa del orden
st.subheader("👀 Vista Previa del Orden Final")
preview_text = " → ".join([f"{idx+1}. {d.name}" for idx, d in enumerate(docs_info)])
st.info(preview_text)

# Sección de combinación
st.header("🔗 Combinar Documentos")

output_name = st.text_input(
    "Nombre del archivo final",
    value=f"documentos_combinados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
    help="El archivo se descargará con este nombre"
)

# Preparar opciones
merge_options = {
    'add_page_break': add_page_breaks,
    'preserve_styles': preserve_styles,
    'add_cover_page': add_cover_page,
    'add_table_of_contents': add_table_of_contents,
    'stop_on_error': stop_on_error,
}

if add_cover_page:
    merge_options['cover_title'] = cover_title
    merge_options['cover_subtitle'] = cover_subtitle
    merge_options['cover_info'] = cover_info

# Barra de progreso
progress_bar = st.progress(0)
status_text = st.empty()

def progress_callback(current, total, message):
    """Callback para actualizar la barra de progreso"""
    progress = current / total
    progress_bar.progress(progress)
    status_text.text(f"{message} ({current}/{total})")

# Botón de combinación
col1, col2 = st.columns([2, 1])

with col1:
    if st.button("🧩 Combinar Documentos", type="primary", use_container_width=True):
        if not docs_info:
            st.error("❌ No hay documentos para combinar")
        else:
            try:
                merger = ProfessionalDocumentMerger(progress_callback=progress_callback)
                
                result_bytes, stats = merger.merge_documents(docs_info, merge_options)
                
                st.session_state.merged_bytes = result_bytes
                st.session_state.merge_stats = stats
                st.session_state.output_name = output_name if output_name.lower().endswith(".docx") else (output_name + ".docx")
                
                progress_bar.progress(1.0)
                status_text.text("✅ ¡Combinación completada!")
                
                st.success("✅ Documentos combinados exitosamente!")
                st.balloons()
                
            except Exception as e:
                progress_bar.empty()
                status_text.empty()
                st.error(f"❌ Error al combinar documentos: {str(e)}")
                with st.expander("Detalles del error"):
                    st.code(traceback.format_exc())
                logger.error(f"Error en combinación: {e}\n{traceback.format_exc()}")

with col2:
    if st.button("🔄 Limpiar Todo", use_container_width=True):
        st.session_state.documents = []
        st.session_state.doc_sources = {}
        st.session_state.merged_bytes = None
        st.session_state.merge_stats = None
        progress_bar.empty()
        status_text.empty()
        st.rerun()

# Sección de descarga
if st.session_state.merged_bytes:
    st.divider()
    st.header("⬇️ Descargar Resultado")
    
    file_size = len(st.session_state.merged_bytes)
    stats = st.session_state.merge_stats or {}
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Tamaño del archivo", format_file_size(file_size))
    with col2:
        st.metric("Tiempo de procesamiento", f"{stats.get('processing_time', 0):.2f}s")
    with col3:
        st.metric("Documentos combinados", stats.get('total_docs', 0))
    
    st.download_button(
        "💾 Descargar Documento Combinado",
        data=st.session_state.merged_bytes,
        file_name=st.session_state.output_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary"
    )
    
    # Mostrar estadísticas detalladas
    with st.expander("📊 Estadísticas Detalladas"):
        st.json(stats)

# Footer con información
with st.expander("ℹ️ Información y Soporte"):
    st.markdown("""
    ### ✅ Características Implementadas
    
    - **Motor Profesional**: Usa `docxcompose` para combinación robusta
    - **Sin Páginas en Blanco**: docxcompose maneja automáticamente los saltos de página
    - **Preservación de Formato**: Estilos, imágenes, tablas se preservan correctamente
    - **Opciones Avanzadas**: Portada, índice, configuración flexible
    
    ### 🔧 Tecnología
    
    Esta aplicación usa **docxcompose**, la librería estándar de la industria para 
    combinar documentos Word. A diferencia de métodos manuales, docxcompose:
    
    - Maneja correctamente las secciones y saltos de página
    - Preserva estilos sin conflictos
    - Evita páginas en blanco innecesarias
    - Es usado en aplicaciones de nivel empresarial
    
    ### ⚠️ Limitaciones Conocidas
    
    - Headers/footers complejos pueden requerir ajuste manual
    - Numeraciones complejas pueden necesitar revisión
    
    ### 💡 Recomendaciones
    
    - Revisa siempre el documento combinado antes de usarlo en producción
    - Guarda copias de los documentos originales
    """)
